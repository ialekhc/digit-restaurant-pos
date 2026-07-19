from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUTPUT = ROOT / "docs" / "Digit_RMS_User_Tutorial_and_Role_Guide.docx"
LOGO = ROOT / "client" / "public" / "digit-nepal" / "mark-light.png"

NAVY = "17324D"
BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
TEAL = "0F766E"
INK = "1F2937"
MUTED = "64748B"
LIGHT_BLUE = "E8EEF5"
LIGHT_TEAL = "ECFDF5"
LIGHT_GOLD = "FFF8E8"
LIGHT_RED = "FEF2F2"
LIGHT_GRAY = "F2F4F7"
WHITE = "FFFFFF"
BORDER = "CBD5E1"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for edge, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths_dxa, indent_dxa=120):
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl_pr = table._tbl.tblPr

    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths_dxa)))
    tbl_w.set(qn("w:type"), "dxa")

    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(indent_dxa))
    tbl_ind.set(qn("w:type"), "dxa")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        row.height = None
        for index, cell in enumerate(row.cells):
            width = widths_dxa[index]
            cell.width = Inches(width / 1440)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            set_cell_margins(cell)
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(width))
            tc_w.set(qn("w:type"), "dxa")


def set_table_borders(table, color=BORDER, size="6"):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        node = borders.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            borders.append(node)
        node.set(qn("w:val"), "single")
        node.set(qn("w:sz"), size)
        node.set(qn("w:color"), color)


def repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    header = OxmlElement("w:tblHeader")
    header.set(qn("w:val"), "true")
    tr_pr.append(header)


def set_run_font(run, name="Calibri", size=None, color=None, bold=None, italic=None):
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def set_paragraph_shading(paragraph, fill, border_color=None):
    p_pr = paragraph._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    p_pr.append(shd)
    if border_color:
        borders = OxmlElement("w:pBdr")
        left = OxmlElement("w:left")
        left.set(qn("w:val"), "single")
        left.set(qn("w:sz"), "18")
        left.set(qn("w:space"), "8")
        left.set(qn("w:color"), border_color)
        borders.append(left)
        p_pr.append(borders)


def add_callout(doc, label, text, tone="info"):
    palette = {
        "info": (LIGHT_BLUE, BLUE),
        "success": (LIGHT_TEAL, TEAL),
        "warning": (LIGHT_GOLD, "B7791F"),
        "danger": (LIGHT_RED, "B91C1C"),
    }
    fill, accent = palette[tone]
    p = doc.add_paragraph(style="Callout")
    set_paragraph_shading(p, fill, accent)
    label_run = p.add_run(f"{label}: ")
    set_run_font(label_run, bold=True, color=accent)
    text_run = p.add_run(text)
    set_run_font(text_run, color=INK)
    return p


def add_numbering_definition(doc, kind):
    numbering = doc.part.numbering_part.element
    abstract_ids = [int(node.get(qn("w:abstractNumId"))) for node in numbering.findall(qn("w:abstractNum"))]
    num_ids = [int(node.get(qn("w:numId"))) for node in numbering.findall(qn("w:num"))]
    abstract_id = max(abstract_ids, default=0) + 1
    num_id = max(num_ids, default=0) + 1

    abstract = OxmlElement("w:abstractNum")
    abstract.set(qn("w:abstractNumId"), str(abstract_id))
    multi = OxmlElement("w:multiLevelType")
    multi.set(qn("w:val"), "singleLevel")
    abstract.append(multi)

    level = OxmlElement("w:lvl")
    level.set(qn("w:ilvl"), "0")
    start = OxmlElement("w:start")
    start.set(qn("w:val"), "1")
    level.append(start)
    num_fmt = OxmlElement("w:numFmt")
    num_fmt.set(qn("w:val"), "bullet" if kind == "bullet" else "decimal")
    level.append(num_fmt)
    level_text = OxmlElement("w:lvlText")
    level_text.set(qn("w:val"), "•" if kind == "bullet" else "%1.")
    level.append(level_text)
    justification = OxmlElement("w:lvlJc")
    justification.set(qn("w:val"), "left")
    level.append(justification)
    p_pr = OxmlElement("w:pPr")
    tabs = OxmlElement("w:tabs")
    tab = OxmlElement("w:tab")
    tab.set(qn("w:val"), "num")
    tab.set(qn("w:pos"), "540")
    tabs.append(tab)
    p_pr.append(tabs)
    ind = OxmlElement("w:ind")
    ind.set(qn("w:left"), "540")
    ind.set(qn("w:hanging"), "270")
    p_pr.append(ind)
    spacing = OxmlElement("w:spacing")
    spacing.set(qn("w:after"), "80")
    spacing.set(qn("w:line"), "300")
    spacing.set(qn("w:lineRule"), "auto")
    p_pr.append(spacing)
    level.append(p_pr)
    if kind == "bullet":
        r_pr = OxmlElement("w:rPr")
        fonts = OxmlElement("w:rFonts")
        fonts.set(qn("w:ascii"), "Symbol")
        fonts.set(qn("w:hAnsi"), "Symbol")
        r_pr.append(fonts)
        level.append(r_pr)
    abstract.append(level)
    numbering.append(abstract)

    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(num_id))
    abstract_ref = OxmlElement("w:abstractNumId")
    abstract_ref.set(qn("w:val"), str(abstract_id))
    num.append(abstract_ref)
    numbering.append(num)
    return num_id


def apply_numbering(paragraph, num_id):
    p_pr = paragraph._p.get_or_add_pPr()
    num_pr = OxmlElement("w:numPr")
    level = OxmlElement("w:ilvl")
    level.set(qn("w:val"), "0")
    num = OxmlElement("w:numId")
    num.set(qn("w:val"), str(num_id))
    num_pr.append(level)
    num_pr.append(num)
    p_pr.append(num_pr)


def add_bullets(doc, items, bullet_num_id):
    for item in items:
        p = doc.add_paragraph(style="List Body")
        apply_numbering(p, bullet_num_id)
        p.add_run(item)


def add_steps(doc, steps, decimal_num_id):
    for title, detail in steps:
        p = doc.add_paragraph(style="List Body")
        apply_numbering(p, decimal_num_id)
        title_run = p.add_run(f"{title}. ")
        title_run.bold = True
        p.add_run(detail)


def add_section_title(doc, title, subtitle=None):
    p = doc.add_paragraph(style="Heading 1")
    p.add_run(title)
    if subtitle:
        s = doc.add_paragraph(style="Section Subtitle")
        s.add_run(subtitle)


def add_role_header(doc, role, purpose, start_screen, role_type="Standard staff role"):
    p = doc.add_paragraph(style="Heading 1")
    p.add_run(role)
    meta = doc.add_paragraph(style="Callout")
    set_paragraph_shading(meta, LIGHT_BLUE, BLUE)
    for index, (label, value) in enumerate((("Primary purpose", purpose), ("Typical start screen", start_screen), ("Availability", role_type))):
        if index:
            meta.add_run("\n")
        set_run_font(meta.add_run(f"{label}: "), bold=True, color=DARK_BLUE)
        set_run_font(meta.add_run(value), color=INK)


def add_role_content(doc, can_do, daily_flow, boundaries, bullet_num_id, decimal_num_id):
    doc.add_paragraph("What this role can do", style="Heading 2")
    add_bullets(doc, can_do, bullet_num_id)
    doc.add_paragraph("Recommended daily workflow", style="Heading 2")
    add_steps(doc, daily_flow, decimal_num_id)
    add_callout(doc, "Permission boundary", boundaries, "warning")


def add_page_break(doc):
    doc.add_page_break()


def add_page_field(paragraph):
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instruction = OxmlElement("w:instrText")
    instruction.set(qn("xml:space"), "preserve")
    instruction.text = " PAGE "
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instruction, separate, text, end])


def configure_styles(doc):
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor.from_string(INK)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    for name, size, color, before, after in (
        ("Heading 1", 16, BLUE, 18, 10),
        ("Heading 2", 13, BLUE, 14, 7),
        ("Heading 3", 12, DARK_BLUE, 10, 5),
    ):
        style = styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    for style_name in ("List Body",):
        try:
            style = styles[style_name]
        except KeyError:
            style = styles.add_style(style_name, 1)
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(11)
        style.paragraph_format.space_after = Pt(4)
        style.paragraph_format.line_spacing = 1.25

    subtitle = styles.add_style("Section Subtitle", 1)
    subtitle.font.name = "Calibri"
    subtitle._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    subtitle._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    subtitle.font.size = Pt(10.5)
    subtitle.font.italic = True
    subtitle.font.color.rgb = RGBColor.from_string(MUTED)
    subtitle.paragraph_format.space_after = Pt(10)

    callout = styles.add_style("Callout", 1)
    callout.font.name = "Calibri"
    callout._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    callout._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    callout.font.size = Pt(10.5)
    callout.paragraph_format.left_indent = Inches(0.12)
    callout.paragraph_format.right_indent = Inches(0.08)
    callout.paragraph_format.space_before = Pt(6)
    callout.paragraph_format.space_after = Pt(9)
    callout.paragraph_format.line_spacing = 1.2


def configure_sections(doc):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.right_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    header = section.header
    p = header.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = p.add_run("DIGIT RMS  |  USER TUTORIAL & ROLE GUIDE")
    set_run_font(run, size=8.5, color=MUTED, bold=True)

    footer = section.footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = p.add_run("Digit Nepal  •  July 2026  •  Page ")
    set_run_font(run, size=8.5, color=MUTED)
    add_page_field(p)


def add_cover(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(34)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if LOGO.exists():
        picture = p.add_run().add_picture(str(LOGO), width=Inches(1.05))
        picture._inline.docPr.set("descr", "Digit Nepal logo")
        picture._inline.docPr.set("title", "Digit Nepal")

    kicker = doc.add_paragraph()
    kicker.alignment = WD_ALIGN_PARAGRAPH.CENTER
    kicker.paragraph_format.space_before = Pt(22)
    kicker.paragraph_format.space_after = Pt(12)
    set_run_font(kicker.add_run("OPERATIONS HANDBOOK"), size=10, color=TEAL, bold=True)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_after = Pt(8)
    set_run_font(title.add_run("Digit Restaurant\nManagement System"), size=30, color=NAVY, bold=True)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.paragraph_format.space_after = Pt(24)
    set_run_font(subtitle.add_run("User Tutorial & Role Guide"), size=17, color=BLUE, bold=True)

    scope = doc.add_paragraph()
    scope.alignment = WD_ALIGN_PARAGRAPH.CENTER
    scope.paragraph_format.space_after = Pt(30)
    set_run_font(scope.add_run("For restaurant owners, administrators, managers, cashiers,\nservice staff, kitchen teams, and specialist users"), size=11.5, color=MUTED)

    add_callout(doc, "Scope", "This handbook intentionally excludes the platform Super Admin portal. It explains restaurant-level work, role boundaries, and the normal order-to-payment process.", "info")

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.paragraph_format.space_before = Pt(28)
    set_run_font(meta.add_run("Production desktop edition  |  Version 1.0  |  July 2026"), size=9.5, color=MUTED, bold=True)


def add_contents(doc, bullet_num_id):
    add_section_title(doc, "How to use this handbook", "Start with the common workflow, then read the section for your assigned role.")
    add_callout(doc, "Access principle", "The menu only shows modules allowed by your role, subscription plan, branch assignment, and any user-specific permission overrides. A missing menu item is usually intentional.", "success")
    doc.add_paragraph("Guide map", style="Heading 2")
    add_bullets(doc, [
        "Getting started: launch, sign in, navigate, change password, and sign out.",
        "Default role map: who is responsible for what and where each role starts.",
        "Core service workflow: table and order creation through kitchen preparation, service, billing, and reporting.",
        "Role playbooks: step-by-step routines for each restaurant-level role.",
        "Module tutorials: users, menu, tables, orders, kitchen stations, billing, inventory, suppliers, reports, settings, and printing.",
        "Shift checklists and troubleshooting: safe opening, handover, closing, and escalation.",
    ], bullet_num_id)
    doc.add_paragraph("Important conventions", style="Heading 2")
    terms = [
        ("Menu label", "The name shown in the left sidebar, such as Orders, Billing, or Inventory."),
        ("Permission", "A server-enforced action such as order.create or payment.collect. Hiding a button does not replace server permission checks."),
        ("Feature", "A module enabled by the restaurant subscription plan. Feature availability can differ by restaurant."),
        ("Branch scope", "The restaurants or branches the signed-in account is allowed to see."),
    ]
    table = doc.add_table(rows=1, cols=2)
    set_table_geometry(table, [2400, 6960])
    set_table_borders(table)
    table.rows[0].cells[0].text = "Term"
    table.rows[0].cells[1].text = "Meaning"
    repeat_table_header(table.rows[0])
    for cell in table.rows[0].cells:
        set_cell_shading(cell, LIGHT_BLUE)
        for run in cell.paragraphs[0].runs:
            set_run_font(run, bold=True, color=DARK_BLUE)
    for term, meaning in terms:
        row = table.add_row()
        row.cells[0].text = term
        row.cells[1].text = meaning
        for run in row.cells[0].paragraphs[0].runs:
            set_run_font(run, bold=True, color=INK)
    set_table_geometry(table, [2400, 6960])


def add_getting_started(doc, decimal_num_id, bullet_num_id):
    add_section_title(doc, "1. Getting started", "Use the production desktop app to work with the hosted restaurant server.")
    doc.add_paragraph("Open and sign in", style="Heading 2")
    add_steps(doc, [
        ("Open the app", "Start Digit Restaurant POS from the desktop shortcut or the installed application."),
        ("Confirm connectivity", "If a connection-status screen appears, check the internet connection and retry. The production app uses the hosted API at https://digitnp.com/api."),
        ("Enter your account", "Use the email and password issued by the restaurant owner or authorized manager. Never use another staff member's account."),
        ("Review your role", "The sidebar displays the signed-in role. Confirm it matches your job before recording orders or payments."),
    ], decimal_num_id)
    add_callout(doc, "Security", "Do not share passwords, save them in public browsers, or leave a signed-in terminal unattended. Each action should remain attributable to the staff member who performed it.", "danger")

    doc.add_paragraph("Navigate the workspace", style="Heading 2")
    add_bullets(doc, [
        "Use the left sidebar to open modules. On a smaller screen, use the menu button in the top bar.",
        "The sidebar groups pages under Operations, Billing, Restaurant Setup, Resources, Insights, and System.",
        "Use search and filters before assuming a record is missing. Orders can be filtered by status, type, or table.",
        "Status dashboards refresh automatically, but a Refresh button is available when immediate confirmation is needed.",
        "Use the top-bar logout action at the end of a shift or whenever another person takes over the terminal.",
    ], bullet_num_id)

    doc.add_paragraph("Change your password", style="Heading 2")
    add_steps(doc, [
        ("Open Settings", "Choose Settings from the System group."),
        ("Enter credentials", "Enter the old password, the new password, and the confirmation."),
        ("Submit", "Select Change Password, then sign in again if requested."),
    ], decimal_num_id)
    add_callout(doc, "If access looks wrong", "Do not borrow another role. Record the missing menu or action, then ask the restaurant owner or authorized manager to review your role, branch assignment, plan features, and permission overrides.", "warning")


def add_role_map(doc):
    add_section_title(doc, "2. Default user-role map", "Super Admin is excluded. These are the default restaurant-level responsibilities.")
    standard = [
        ("Restaurant Owner", "Full restaurant control", "Dashboard"),
        ("Admin", "Broad operations and configuration", "Dashboard"),
        ("Manager", "Day-to-day branch operations", "Orders"),
        ("Cashier", "Billing, payments, receipts, own shift", "Billing"),
        ("Waiter", "Tables, order entry, service updates", "Create Order"),
        ("Kitchen", "Food queue and preparation updates", "Kitchen Display"),
        ("Barista", "Bar queue and drink preparation", "Bar Display"),
        ("Customer", "Optional authenticated ordering profile", "Orders"),
    ]
    table = doc.add_table(rows=1, cols=3)
    set_table_geometry(table, [2200, 4680, 2480])
    set_table_borders(table)
    for index, label in enumerate(("Standard role", "Primary responsibility", "Default start")):
        table.rows[0].cells[index].text = label
        set_cell_shading(table.rows[0].cells[index], LIGHT_BLUE)
        for run in table.rows[0].cells[index].paragraphs[0].runs:
            set_run_font(run, bold=True, color=DARK_BLUE)
    repeat_table_header(table.rows[0])
    for role, responsibility, start in standard:
        row = table.add_row()
        for index, value in enumerate((role, responsibility, start)):
            row.cells[index].text = value
        for run in row.cells[0].paragraphs[0].runs:
            set_run_font(run, bold=True, color=INK)
    set_table_geometry(table, [2200, 4680, 2480])

    doc.add_paragraph("Specialized permission profiles", style="Heading 2")
    specialized = [
        ("Chef", "Same core kitchen preparation profile as Kitchen."),
        ("Inventory Manager", "Stock, transfers, wastage, purchases, and suppliers."),
        ("Accountant", "Sales/profit reports, exports, payment viewing, and expenses."),
        ("Delivery Partner", "Assigned-order viewing and delivery status updates."),
        ("Customer Support", "Read-only order and customer lookup."),
    ]
    table = doc.add_table(rows=1, cols=2)
    set_table_geometry(table, [2800, 6560])
    set_table_borders(table)
    table.rows[0].cells[0].text = "Specialized role"
    table.rows[0].cells[1].text = "Default purpose"
    for cell in table.rows[0].cells:
        set_cell_shading(cell, LIGHT_GRAY)
        for run in cell.paragraphs[0].runs:
            set_run_font(run, bold=True, color=DARK_BLUE)
    repeat_table_header(table.rows[0])
    for role, purpose in specialized:
        row = table.add_row()
        row.cells[0].text = role
        row.cells[1].text = purpose
    set_table_geometry(table, [2800, 6560])
    add_callout(doc, "Provisioning note", "The current Users screen offers Restaurant Owner, Admin, Manager, Cashier, Waiter, Kitchen, Barista, and Customer. Specialized roles are supported by the permission model but may require administrator/API provisioning until they are enabled in the standard form.", "warning")
    add_callout(doc, "Defaults can change", "The server can add or deny permissions for an individual account. Subscription features can also hide modules. Treat this guide as the default role baseline, not a substitute for the access shown in your own account.", "info")


def add_core_workflow(doc, decimal_num_id, bullet_num_id):
    add_section_title(doc, "3. Core restaurant workflow", "The normal path from guest seating to completed payment.")
    add_steps(doc, [
        ("Prepare the table and menu", "Owner/Admin/Manager confirms tables, menu items, availability, prices, printers, and receipt content."),
        ("Create the order", "Waiter, Cashier, Manager, or Barista selects Dine In or Takeaway, chooses a table/customer when relevant, adds items and notes, then submits the order."),
        ("Route production", "Food, drink, and smoke items are routed to their configured station displays and print queues."),
        ("Prepare items", "Kitchen or Bar staff marks the order Preparing, then records each item as Ready. Service staff records each ready item as Served; takeaway items are shown as Packed."),
        ("Collect payment", "Cashier opens Billing, looks up the table, verifies READY or SERVED unpaid orders, selects payment method/status, records discount and amount paid, then accepts payment."),
        ("Issue receipt and review", "Cashier prints or downloads the receipt. Managers and owners review reports, exceptions, low stock, and shift outcomes."),
    ], decimal_num_id)

    doc.add_paragraph("Order status meaning", style="Heading 2")
    statuses = [
        ("PENDING", "Submitted and waiting for a production station to accept/start."),
        ("PREPARING", "The kitchen/bar/smoke station is working on the items."),
        ("READY", "All applicable items are ready for service or packing."),
        ("SERVED", "Dine-in items have been delivered; takeaway items are displayed as packed."),
        ("COMPLETED", "The service and payment lifecycle is finished."),
        ("CANCELLED", "The order was cancelled by an authorized role with a reason."),
    ]
    table = doc.add_table(rows=1, cols=2)
    set_table_geometry(table, [1900, 7460])
    set_table_borders(table)
    table.rows[0].cells[0].text = "Status"
    table.rows[0].cells[1].text = "Operational meaning"
    for cell in table.rows[0].cells:
        set_cell_shading(cell, LIGHT_BLUE)
        for run in cell.paragraphs[0].runs:
            set_run_font(run, bold=True, color=DARK_BLUE)
    repeat_table_header(table.rows[0])
    for status, meaning in statuses:
        row = table.add_row()
        row.cells[0].text = status
        row.cells[1].text = meaning
        for run in row.cells[0].paragraphs[0].runs:
            set_run_font(run, bold=True)
    set_table_geometry(table, [1900, 7460])
    add_callout(doc, "Billing rule", "Billing combines billable unpaid orders for the selected table. In the current screen, billable orders are READY or SERVED. Confirm the table and item quantities before accepting payment.", "warning")
    doc.add_paragraph("Operational controls", style="Heading 2")
    add_bullets(doc, [
        "Use item notes for preparation details, allergies, or service instructions; confirm critical allergy information verbally as required by restaurant policy.",
        "Print station tickets only to the intended printer route: Food to Kitchen; Bar, Smoke, bills, and receipts to Reception unless configured otherwise.",
        "For table transfers or split/merge actions, verify both source and destination before confirming.",
        "Never delete, cancel, discount, refund, or void a transaction merely to correct a training mistake; follow the restaurant's authorization and audit procedure.",
    ], bullet_num_id)


def add_owner_admin_manager(doc, bullet_num_id, decimal_num_id):
    add_role_header(doc, "4. Restaurant Owner", "Restaurant-wide setup, access, operations, reporting, and oversight", "Dashboard")
    add_role_content(doc, [
        "View and update restaurant/branch operations, menus, tables, orders, billing, inventory, suppliers, customers, reports, settings, and audit-related data.",
        "Create and manage staff accounts within the restaurant and assign supported subordinate roles/branches.",
        "Approve high-impact operational actions within configured approval limits.",
        "Review subscription information and the enabled feature set for the restaurant.",
    ], [
        ("Review Dashboard", "Check sales collected, active order load, table occupancy, cancellations, best sellers, and low-stock alerts."),
        ("Check staffing", "Open My Users and confirm active accounts, correct roles, and branch scope."),
        ("Review setup", "Confirm menu availability, table status, suppliers, inventory, receipt content, and printer routes."),
        ("Monitor operations", "Use Orders, station displays, Billing, and Reports to identify exceptions."),
        ("Close the day", "Review reports and exports, investigate cancellations/discounts, and confirm shift handover."),
    ], "Owner access is restaurant-scoped, not platform-wide. It does not include the Super Admin vendor, plan, or global subscription portal.", bullet_num_id, decimal_num_id)

    doc.add_paragraph("Creating a staff account", style="Heading 2")
    add_steps(doc, [
        ("Open My Users", "Select Users from the System group."),
        ("Enter staff details", "Provide name, unique email, temporary password, optional phone, role, and Active status."),
        ("Choose least privilege", "Select only the role needed for the person's work. Do not use Admin for convenience."),
        ("Create and test", "Create the account, have the staff member sign in, and confirm only the expected modules are visible."),
        ("Secure the handoff", "Ask the staff member to change the temporary password in Settings."),
    ], decimal_num_id)

    doc.add_paragraph("Restaurant Admin", style="Heading 1")
    add_role_content(doc, [
        "Use the Dashboard and operate most restaurant modules: menu, tables, orders, payments, inventory, suppliers, customers, reports, settings, and stations.",
        "Configure operational data and resolve routine issues across the restaurant.",
        "Apply discounts/refunds/voids only within configured limits and local policy.",
    ], [
        ("Open Dashboard", "Review the operation snapshot and low-stock alerts."),
        ("Validate setup", "Check menu availability, tables, inventory, suppliers, and printer configuration."),
        ("Support the floor", "Monitor Orders, kitchen/bar/smoke displays, and Billing for stuck transactions."),
        ("Review reports", "Export or review allowed reports and investigate exceptions."),
    ], "By default, Admin does not manage Users and cannot perform owner-only destructive actions such as order deletion or bill deletion. Escalate account changes to the Restaurant Owner.", bullet_num_id, decimal_num_id)

    doc.add_paragraph("Manager", style="Heading 1")
    add_role_content(doc, [
        "Manage orders, tables, payments, split payments, kitchen visibility, customers, inventory adjustments/wastage, purchases, suppliers, and branch reports.",
        "Create or update subordinate staff where server hierarchy and user-management rules allow; cannot assign peer or higher roles.",
        "Export branch reports and monitor the manager's own shift activity.",
    ], [
        ("Start in Orders", "Review active, ready, served, and cancelled orders; filter by table when resolving a guest issue."),
        ("Check production", "Inspect station displays and prioritize ready dishes before new work."),
        ("Control service", "Resolve table transfers, split/merge requests, cancellations, and discount requests within permission."),
        ("Check stock", "Review low stock, record approved adjustments/wastage, and create purchase movements."),
        ("Handover", "Export the branch report and document unresolved orders, payments, and stock issues."),
    ], "Manager has a limited-dashboard permission but the current navigation starts at Orders. Some deployments may not expose the Dashboard page to this role. Higher-role assignment and owner-level settings remain restricted.", bullet_num_id, decimal_num_id)


def add_waiter(doc, bullet_num_id, decimal_num_id):
    add_role_header(doc, "5. Waiter / Service Staff", "Table service, order entry, station visibility, and serving updates", "Create Order")
    add_role_content(doc, [
        "View and assign tables; create dine-in and takeaway orders.",
        "Add more items to an active table order, add preparation notes, and transfer permitted orders/tables.",
        "View kitchen, bar, and smoke queues and record ready items as Served when permitted.",
        "View/create/update customer records and review own-order activity.",
    ], [
        ("Check tables", "Open Tables and confirm the guest's table is available or occupied as expected."),
        ("Create the order", "Select order type, table, and optional customer; add items by Food, Drink, or Smoke category."),
        ("Add details", "Set quantity and use Notes for preparation instructions. Review the Order Summary."),
        ("Submit", "Create the order and confirm it appears in Orders and the correct station display."),
        ("Serve", "When an item is ready, confirm the table/order and record the item as Served. For takeaway, confirm it is Packed."),
        ("Add rounds", "From the selected table's order view, choose New Order For This Table and submit only the additional items."),
    ], "Waiter cannot collect payment, apply discounts by default, manage menu/inventory, cancel orders, or modify users. Send billing to Cashier and exceptions to Manager.", bullet_num_id, decimal_num_id)

    doc.add_paragraph("Order-entry checklist", style="Heading 2")
    add_bullets(doc, [
        "Correct order type: DINE_IN or TAKEAWAY.",
        "Correct table and customer, if used.",
        "Correct menu tab, category, item, quantity, price, and notes.",
        "No duplicate items caused by double-clicking or re-submitting.",
        "Order visible in Orders and routed to the intended station.",
    ], bullet_num_id)
    add_callout(doc, "Guest changes", "If production has already started, do not silently replace items. Follow the restaurant's cancellation/void procedure and involve the Manager when authorization is required.", "warning")


def add_kitchen_bar(doc, bullet_num_id, decimal_num_id):
    add_role_header(doc, "6. Kitchen / Chef", "Food preparation queue, item readiness, and menu availability", "Kitchen Display", "Kitchen is standard; Chef is a specialized equivalent profile")
    add_role_content(doc, [
        "View new, preparing, ready, and recently completed kitchen tickets.",
        "Mark a pending order Preparing, then record readiness one item at a time with +1 Ready.",
        "Mark menu items unavailable when the station cannot produce them, if the action is enabled.",
        "See table/order number, order type, item quantity, notes, age, and readiness/served counts.",
    ], [
        ("Review Ready first", "The display intentionally shows ready dishes before preparing and new orders. Clear ready items quickly."),
        ("Start work", "Confirm the ticket and choose Mark Preparing."),
        ("Prepare by item", "Read notes, prepare the required quantity, and choose +1 Ready for each completed unit."),
        ("Confirm completion", "Ensure ready quantity matches ordered quantity. Recently completed tickets remain visible for quick verification."),
        ("Report shortages", "Mark an item unavailable where permitted and immediately notify Manager/service staff."),
    ], "Do not collect payment, alter customers, or change unrelated order details. Kitchen users should not mark Served unless the restaurant explicitly assigns that responsibility; service staff normally confirms delivery.", bullet_num_id, decimal_num_id)

    doc.add_paragraph("Barista", style="Heading 1")
    add_role_content(doc, [
        "Operate the Bar Display, accept/prepare drink tickets, and update item readiness.",
        "View/create/update operational orders and send items to production where enabled.",
        "View menu availability and create customer records when required for service.",
    ], [
        ("Open Bar Display", "Confirm the station label and review ready, preparing, and recent new drink orders."),
        ("Start and prepare", "Mark the ticket Preparing and use +1 Ready per finished drink."),
        ("Coordinate handoff", "Keep order/table identifiers attached to the drinks and notify service staff when complete."),
        ("Update availability", "If an ingredient or item is unavailable, update availability where permitted and notify Manager."),
    ], "Barista is not a Cashier role. Payment collection, refunds, and register work remain restricted unless an explicit permission override is granted.", bullet_num_id, decimal_num_id)

    doc.add_paragraph("Smoke Display", style="Heading 2")
    add_bullets(doc, [
        "The Smoke Display follows the same station workflow: New -> Preparing -> Ready -> Served/Completed.",
        "Station visibility uses the kitchen-view permission, so the exact staff role operating it depends on local policy.",
        "Reception printing handles Bar and Smoke tickets in the default printer-routing model.",
    ], bullet_num_id)


def add_cashier(doc, bullet_num_id, decimal_num_id):
    add_role_header(doc, "7. Cashier", "Billing, payment collection, receipt handling, and own-shift reporting", "Billing")
    add_role_content(doc, [
        "View/create/update orders needed for billing, collect payments, and split payments when enabled.",
        "Open/view/close the assigned cash register and review own-shift/own-order information.",
        "Apply limited discounts within the account's configured percentage limit.",
        "Create/update customers and print or reprint receipts.",
    ], [
        ("Confirm readiness", "Verify the correct table's unpaid orders are READY or SERVED before billing."),
        ("Look up the table", "Open Billing and enter the exact table number."),
        ("Review the bill", "Check order count, items, quantities, subtotal, and any existing payments."),
        ("Enter payment", "Choose Cash, Card, QR, Online, or Split as enabled; choose status; enter discount and amount paid."),
        ("Check change", "For cash payments, confirm the automatically calculated change before accepting."),
        ("Accept and issue receipt", "Select Accept Payment, then print or download the receipt. Return change only after the payment is confirmed."),
        ("Close shift", "Reconcile the shift, report discrepancies, and close the register according to restaurant policy."),
    ], "Cashier cannot manage menu, stock, suppliers, staff, or broad restaurant settings by default. Refunds/voids are not part of the default Cashier permission set even though monetary approval limits exist; escalate them.", bullet_num_id, decimal_num_id)

    doc.add_paragraph("Billing controls", style="Heading 2")
    add_bullets(doc, [
        "Never accept payment against the wrong table. Confirm table number verbally with service staff when uncertain.",
        "Use PAID only when the full required amount is received; use PARTIAL only under approved restaurant policy.",
        "If a discount exceeds your limit, stop and obtain Manager/Owner authorization instead of splitting or editing the bill to bypass controls.",
        "Use Bills & Payment History to search by bill, order, table, or cashier and to reprint/download the correct receipt.",
        "Export to Excel only when authorized; exported files may contain sensitive sales and customer data.",
    ], bullet_num_id)
    add_callout(doc, "Register note", "Cash-register pages exist but are currently hidden from sidebar navigation. Use the restaurant's approved register workflow until the feature is formally enabled.", "info")


def add_specialists(doc, bullet_num_id, decimal_num_id):
    add_section_title(doc, "8. Specialist and customer profiles", "These profiles are server-supported but may require administrator provisioning.")

    roles = [
        ("Inventory Manager", [
            "View/manage inventory; record adjustments, transfers, and wastage.",
            "Create and approve purchase movements; manage supplier records.",
        ], "Inventory, Purchase In/Out, Suppliers", "Cannot collect payment, manage orders, or access restaurant-wide settings by default."),
        ("Accountant", [
            "View branch and restaurant sales, profit information, and allowed exports.",
            "View payments and expenses for reconciliation and reporting.",
        ], "Reports", "Read/analysis profile; cannot alter orders, collect payments, or manage staff by default."),
        ("Delivery Partner", [
            "View assigned/accessible orders and update delivery status.",
        ], "Orders", "Access should remain limited to assigned delivery work; no payment, customer-management, or restaurant-setup rights by default."),
        ("Customer Support", [
            "Look up order and customer information to answer service questions.",
        ], "Orders / Customers", "Read-oriented profile; do not edit transactions or disclose customer/order information without identity verification."),
        ("Customer", [
            "Optional authenticated profile for viewing, creating, updating, or cancelling own permitted orders and viewing payment state.",
        ], "Orders", "The customer role is optional and not enabled in standard frontend routing by default. Public table QR ordering is the normal guest-facing path."),
    ]

    for role, abilities, start, boundary in roles:
        doc.add_paragraph(role, style="Heading 2")
        p = doc.add_paragraph()
        label = p.add_run("Primary workspace: ")
        label.bold = True
        p.add_run(start)
        add_bullets(doc, abilities, bullet_num_id)
        add_callout(doc, "Boundary", boundary, "warning")

    doc.add_paragraph("Inventory movement procedure", style="Heading 2")
    add_steps(doc, [
        ("Select movement", "Open Purchase In for stock received or Purchase Out for approved stock removal."),
        ("Choose item", "Select the inventory item and confirm its unit."),
        ("Enter transaction", "Record quantity, rate, amount, payment mode, invoice/bill number, transaction date, and notes."),
        ("Review", "Confirm the movement appears in Daily Entry History and that the inventory balance is reasonable."),
        ("Investigate exceptions", "Do not create a second movement to hide an error. Escalate and use the approved correction procedure."),
    ], decimal_num_id)


def add_module_tutorials(doc, bullet_num_id, decimal_num_id):
    add_section_title(doc, "9. Module tutorials", "Use only the modules and actions visible to your own account.")

    doc.add_paragraph("Tables", style="Heading 2")
    add_bullets(doc, [
        "Owners/Admins/authorized Managers can add or edit table number, seating capacity, and status.",
        "Service staff can open a table to view active orders, start another order, and perform permitted transfers.",
        "Permanent table QR codes connect guest ordering to the same dine-in pipeline used by staff and production stations.",
        "Before transferring a table, verify source, destination, item quantities, and active orders.",
    ], bullet_num_id)

    doc.add_paragraph("Menu and availability", style="Heading 2")
    add_bullets(doc, [
        "Menu Categories organizes products; Food, Drink, and Smoke item pages represent separate service/production types.",
        "Use consistent names, categories, prices, preparation station, and availability.",
        "Mark unavailable items promptly so service staff cannot continue selecting them; restore availability only after confirmation.",
    ], bullet_num_id)

    doc.add_paragraph("Inventory and suppliers", style="Heading 2")
    add_steps(doc, [
        ("Create supplier", "Enter the supplier name and relevant phone, email, company, and address."),
        ("Create inventory item", "Enter name, category, quantity, unit, minimum level, supplier, purchase price, and expiry date."),
        ("Monitor", "Use category/time filters, expiring-stock cards, and the low-stock list."),
        ("Record movements", "Use Purchase In/Out or authorized adjustment/wastage actions; include traceable notes and invoice references."),
        ("Reconcile", "Investigate unexpected negative, duplicate, expired, or unusually large movements."),
    ], decimal_num_id)

    doc.add_paragraph("Customers", style="Heading 2")
    add_bullets(doc, [
        "Create or update name, phone, email, address, and loyalty points only when authorized and accurate.",
        "Use Order History to answer service questions, but disclose data only after appropriate identity verification.",
        "Avoid duplicate customers: search by phone/email before creating a new record.",
    ], bullet_num_id)

    doc.add_paragraph("Reports", style="Heading 2")
    add_bullets(doc, [
        "Select daily, weekly, monthly, or yearly period as available, then Refresh Reports.",
        "Review sales totals, best-selling item distribution/table, and low-stock report.",
        "Export only the data needed for an authorized business purpose; store exported files securely.",
    ], bullet_num_id)

    doc.add_paragraph("Settings and printing", style="Heading 2")
    add_bullets(doc, [
        "Receipt Content controls the restaurant name, phone, address, email, and footer printed on bills and tickets on this device.",
        "Printer Settings maps logical printer purpose to the exact OS/QZ Tray printer name, connection details, paper width, and copies.",
        "Print Station claims pending jobs to prevent duplicates, routes them to configured printers, and records success/failure.",
        "Use test print before live service. Do not repeatedly click Print Now while a printer is slow; check Recent Print Results first.",
    ], bullet_num_id)
    add_callout(doc, "Device-specific settings", "Receipt and some printing preferences are stored on the current device. Configure and test every production terminal separately.", "info")


def add_checklists_troubleshooting(doc, bullet_num_id, decimal_num_id):
    add_section_title(doc, "10. Shift checklists & troubleshooting", "Consistent handover protects sales, stock, and guest service.")

    doc.add_paragraph("Opening checklist", style="Heading 2")
    add_bullets(doc, [
        "Internet connection is working and the desktop app opens without a connection warning.",
        "You are signed in with your own account and correct role/branch.",
        "Tables, menu availability, prices, station displays, and printers are correct.",
        "Cashier opening/register procedure is complete under local policy.",
        "Low stock, expiring items, pending orders, and unresolved previous-shift issues are reviewed.",
    ], bullet_num_id)

    doc.add_paragraph("Closing / handover checklist", style="Heading 2")
    add_bullets(doc, [
        "No unexplained PENDING, PREPARING, READY, SERVED, PARTIAL, or UNPAID transactions remain.",
        "Receipts, payment methods, discounts, cancellations, and cash totals are reconciled.",
        "Stock movements, wastage, and shortages are recorded with traceable reasons.",
        "Printer failures and unprinted jobs are resolved or documented.",
        "Reports/exports are stored securely, unresolved issues are handed over, and every terminal is logged out.",
    ], bullet_num_id)

    doc.add_paragraph("Troubleshooting sequence", style="Heading 2")
    add_steps(doc, [
        ("Read the message", "Capture the exact on-screen error, module, record number, time, and action attempted."),
        ("Check scope", "Confirm role, branch, plan feature, and whether the action normally belongs to another role."),
        ("Refresh safely", "Use the page Refresh action once. Do not repeatedly submit an order or payment."),
        ("Check connectivity", "If several pages fail, confirm internet access and whether the hosted service is reachable."),
        ("Escalate with evidence", "Provide the order/table/bill number and a screenshot where policy permits. Never send passwords or full payment details."),
    ], decimal_num_id)

    incidents = [
        ("Menu item missing", "Check menu type/category/availability and plan feature", "Manager or Admin"),
        ("Cannot access a module", "Confirm role, branch, feature, and override", "Owner / authorized Manager"),
        ("Order not on station display", "Confirm station routing and order status; refresh once", "Manager / Admin"),
        ("Payment may be duplicated", "Stop; search payment history by table/order/bill", "Manager / Owner"),
        ("Printer failed", "Check QZ/OS printer name, queue, and Recent Print Results", "Admin / device support"),
        ("Stock balance unexpected", "Review purchase/adjustment/wastage history", "Inventory Manager / Owner"),
        ("Customer data concern", "Limit disclosure and preserve audit evidence", "Owner / privacy contact"),
    ]
    table = doc.add_table(rows=1, cols=3)
    set_table_geometry(table, [2550, 4650, 2160])
    set_table_borders(table)
    for index, label in enumerate(("Issue", "First check", "Escalate to")):
        table.rows[0].cells[index].text = label
        set_cell_shading(table.rows[0].cells[index], LIGHT_BLUE)
        for run in table.rows[0].cells[index].paragraphs[0].runs:
            set_run_font(run, bold=True, color=DARK_BLUE)
    repeat_table_header(table.rows[0])
    for issue, check, owner in incidents:
        row = table.add_row()
        for index, value in enumerate((issue, check, owner)):
            row.cells[index].text = value
    set_table_geometry(table, [2550, 4650, 2160])

    add_callout(doc, "Production safety", "Never test with fake payments, real guest accounts, or destructive actions in production. Use an approved test environment or a clearly authorized training record.", "danger")


def add_quick_reference(doc):
    add_section_title(doc, "Quick reference", "A concise handover page for production terminals.")
    rows = [
        ("Waiter", "Create/extend order; confirm station routing", "Cashier for payment; Manager for exceptions"),
        ("Kitchen/Chef", "Preparing -> +1 Ready per item", "Manager for shortage or incorrect ticket"),
        ("Barista", "Prepare drink queue; mark ready", "Manager; Cashier for payment"),
        ("Cashier", "Verify table -> accept payment -> receipt", "Manager/Owner for refund, void, excess discount"),
        ("Manager", "Monitor floor, payments, staff, stock, reports", "Owner for policy/account/owner-only actions"),
        ("Admin", "Configuration and broad operational support", "Owner for Users and owner-only actions"),
        ("Owner", "Restaurant oversight, users, reporting, approvals", "Platform support for hosted-service issues"),
    ]
    table = doc.add_table(rows=1, cols=3)
    set_table_geometry(table, [1900, 4320, 3140])
    set_table_borders(table)
    for index, label in enumerate(("Role", "Normal action", "Escalation")):
        table.rows[0].cells[index].text = label
        set_cell_shading(table.rows[0].cells[index], NAVY)
        for run in table.rows[0].cells[index].paragraphs[0].runs:
            set_run_font(run, bold=True, color=WHITE)
    repeat_table_header(table.rows[0])
    for role, action, escalation in rows:
        row = table.add_row()
        for index, value in enumerate((role, action, escalation)):
            row.cells[index].text = value
        for run in row.cells[0].paragraphs[0].runs:
            set_run_font(run, bold=True)
    set_table_geometry(table, [1900, 4320, 3140])

    add_callout(doc, "Golden rule", "Use your own account, verify the table/order/bill before acting, and stop to escalate whenever the requested action exceeds your role or approval limit.", "success")

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(24)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_run_font(p.add_run("Digit Restaurant Management System"), size=15, color=NAVY, bold=True)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_run_font(p.add_run("Production user tutorial • Super Admin excluded • July 2026"), size=9.5, color=MUTED)


def build_document():
    doc = Document()
    configure_styles(doc)
    configure_sections(doc)
    bullet_num_id = add_numbering_definition(doc, "bullet")
    decimal_num_id = add_numbering_definition(doc, "decimal")

    add_cover(doc)
    add_page_break(doc)
    add_contents(doc, bullet_num_id)
    add_page_break(doc)
    add_getting_started(doc, decimal_num_id, bullet_num_id)
    add_page_break(doc)
    add_role_map(doc)
    add_page_break(doc)
    add_core_workflow(doc, decimal_num_id, bullet_num_id)
    add_page_break(doc)
    add_owner_admin_manager(doc, bullet_num_id, decimal_num_id)
    add_page_break(doc)
    add_waiter(doc, bullet_num_id, decimal_num_id)
    add_page_break(doc)
    add_kitchen_bar(doc, bullet_num_id, decimal_num_id)
    add_page_break(doc)
    add_cashier(doc, bullet_num_id, decimal_num_id)
    add_page_break(doc)
    add_specialists(doc, bullet_num_id, decimal_num_id)
    add_page_break(doc)
    add_module_tutorials(doc, bullet_num_id, decimal_num_id)
    add_page_break(doc)
    add_checklists_troubleshooting(doc, bullet_num_id, decimal_num_id)
    add_page_break(doc)
    add_quick_reference(doc)

    doc.core_properties.title = "Digit RMS User Tutorial and Role Guide"
    doc.core_properties.subject = "Restaurant system training and default role responsibilities"
    doc.core_properties.author = "Digit Nepal"
    doc.core_properties.keywords = "Digit RMS, restaurant POS, user tutorial, roles, operations"
    doc.core_properties.comments = "Super Admin intentionally excluded."

    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    build_document()
